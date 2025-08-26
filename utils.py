import os
import docx
from lxml import etree
from pinecone import Pinecone
import streamlit as st
from langchain_openai import OpenAIEmbeddings
from langchain.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_pinecone import PineconeVectorStore

# --- Pinecone 連線與操作 ---

@st.cache_resource
def get_pinecone_client():
    """快取 Pinecone 連線，避免重複初始化。"""
    return Pinecone(api_key=os.environ.get("PINECONE_API_KEY"))

def fetch_pinecone_namespaces(index_name):
    """從 Pinecone 獲取所有已存在的 Namespaces 列表。"""
    pc = get_pinecone_client()
    try:
        index_stats = pc.describe_index(index_name).stats
        return list(index_stats.namespaces.keys()) if index_stats and index_stats.namespaces else []
    except Exception:
        return []

# --- GCO 經驗提取函式 (來自 Extract_GCO_wisdom.py) ---

def extract_revisions_from_single_doc(file_path, nsmap):
    """從單一 .docx 檔案中，直接讀取追蹤修訂的內容。"""
    doc = docx.Document(file_path)
    extracted_data = []
    for para in doc.paragraphs:
        if '<w:ins' in para._p.xml or '<w:del' in para._p.xml:
            original_text, revised_text = "", ""
            p_tree = etree.fromstring(para._p.xml)
            runs = p_tree.xpath('.//w:r', namespaces=nsmap)
            for run in runs:
                text_nodes = run.xpath('.//w:t', namespaces=nsmap)
                text = text_nodes[0].text if text_nodes and text_nodes[0].text else ""
                if run.xpath('.//w:ins', namespaces=nsmap): revised_text += text
                elif run.xpath('.//w:del', namespaces=nsmap): original_text += text
                else:
                    original_text += text
                    revised_text += text
            if original_text.strip() != revised_text.strip():
                wisdom_chunk = (
                    f"【審閱案例 - 追蹤修訂】\n"
                    f"  - 修訂前原文：\n"
                    f"    ---\n"
                    f"    {original_text.strip()}\n"
                    f"    ---\n"
                    f"  - 修訂後建議：\n"
                    f"    ---\n"
                    f"    {revised_text.strip()}\n"
                    f"    ---\n"
                )
                extracted_data.append(wisdom_chunk)
    return extracted_data


def extract_comments_from_docx(file_path):
    """從 .docx 檔案中提取所有註解及其關聯的文字。"""
    doc = docx.Document(file_path)
    extracted_data = []
    for comment in doc.comments:
        paragraphs = comment.paragraphs
        original_text = "\n".join([p.text for p in paragraphs]).strip()
        comment_text = comment.text.strip()
        if original_text and comment_text:
            wisdom_chunk = (
                f"【審閱案例 - 法務專家註解】\n"
                f"  - 關聯原文：\n"
                f"    ---\n"
                f"    {original_text.strip()}\n"
                f"    ---\n"
                f"  - 專家註解：\n"
                f"    ---\n"
                f"    {comment_text}\n"
                f"    ---\n"
            )
            extracted_data.append(wisdom_chunk)
    return extracted_data


# --- 資料上傳 (Ingest) 函式 ---

def ingest_docs_to_pinecone(docs, index_name, namespace):
    """將 LangChain 文件區塊上傳至指定的 Pinecone Namespace。"""
    embeddings = OpenAIEmbeddings(model="text-embedding-3-small")
    PineconeVectorStore.from_documents(
        docs,
        embedding=embeddings,
        index_name=index_name,
        namespace=namespace
    )
